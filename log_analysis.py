from docx import Document
from docx.shared import Mm

import postgres_init
import configs


def unic_user():
    """
    Функция подключается к таблице, где содержаться количество пользователей.
    Выгружает количество записей и считает их уникальный список
    :return: количество уникальных пользователей которые имеют доступ
    """
    try:
        access = "Доступ запрещен"
        cursor, conn = postgres_init.postgres_init()
        cursor.execute('SELECT DISTINCT user_id FROM initialization_user WHERE access <> %s', (access,))
        list_unic_user = cursor.fetchall()
        cl_unic_user = len(list_unic_user)
        return cl_unic_user
    except (Exception,):
        return "Не определенно"


def list_of_departments():
    cursor, conn = postgres_init.postgres_init()
    # cursor.execute("SELECT DISTINCT department FROM initialization_user")
    # departments = cursor.fetchall()

    cursor.execute('SELECT department FROM initialization_user')
    departments = cursor.fetchall()

    # Создание словаря с уникальными значениями и их повторениями
    department_count = {}
    for department in departments:
        department = department[0]
        if department in department_count:
            department_count[department] += 1
        else:
            department_count[department] = 1

    return department_count


def main_log_analysis():
    cursor, conn = postgres_init.postgres_init()
    cursor.execute('SELECT * FROM initialization_user')
    table_unic_user = cursor.fetchall()

    '''
    Создаем документ, и настраиваем стили
    '''
    document = Document()
    # доступ к первой секции:
    section = document.sections[0]
    # левое поле в миллиметрах
    section.left_margin = Mm(20)
    # правое поле в миллиметрах
    section.right_margin = Mm(20)
    # верхнее поле в миллиметрах
    section.top_margin = Mm(15)
    # нижнее поле в миллиметрах
    section.bottom_margin = Mm(10)
    document.add_heading("Анализ лога системы", 0)
    # document.add_paragraph(bank_list)

    # количество уникальный пользователей
    cl_unic_user = unic_user()
    document.add_paragraph(f"Количество активных пользователей системы: {cl_unic_user}")
    document.add_paragraph("Полный список пользователей системы: ")

    # создаем таблицу, применяем стили
    table_2 = document.add_table(rows=1, cols=5, style='Light Shading Accent 1')
    table_2.autofit = True
    hdr_cells = table_2.rows[0].cells

    hdr_cells[0].text = "ID пользователя"
    hdr_cells[1].text = "Имя пользователя"
    hdr_cells[2].text = "Username пользователя"
    hdr_cells[3].text = "Отдел пользователя"
    hdr_cells[4].text = "Уровень доступа"

    # убираем guid из hash таблицы
    new_list = []
    for i in table_unic_user:
        new_list_el = list(i)
        new_list_el.pop(4)
        new_list.append(new_list_el)

    # добавляем данные в таблицу
    for id_user, name, username, depart, access in new_list:
        row_cells = table_2.add_row().cells
        row_cells[0].text = str(id_user)
        row_cells[1].text = str(name)
        row_cells[2].text = str(username)
        row_cells[3].text = str(depart)
        row_cells[4].text = str(access)

    # создаем таблицу уникальных отделов и их количество
    document.add_paragraph(' ')
    document.add_paragraph(f'Список используемых отделов в БД RFC Informer Bot: ')

    # получаем из БД отделы и количество пользователей в ней
    department_count = list_of_departments()

    # создаем таблицу, применяем стили
    table_2 = document.add_table(rows=1, cols=2, style='Light Shading Accent 1')
    table_2.autofit = True
    hdr_cells = table_2.rows[0].cells

    hdr_cells[0].text = "Добавленные отделы"
    hdr_cells[1].text = "Количество пользователей"

    # добавляем данные в таблицу
    for departments, cl_users in department_count.items():
        row_cells = table_2.add_row().cells
        row_cells[0].text = str(departments)
        row_cells[1].text = str(cl_users)

    document.save(configs.log_analysis_filename)