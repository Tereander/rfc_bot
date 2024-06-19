import traceback
import time

from docx import Document
from docx.shared import Inches
from datetime import datetime

import logs
import firewall_mars
import postgres_init
import configs


def create_file_error(error, error_basic):
    """
    Функция принимает на вход текст ошибки и создает файл для отчетности
    param error:: Полный лог ошибки
    param error_basic:: Краткий лог ошибки
    :return:
    """
    time_now = str(datetime.now())
    date_src = time_now[:19]

    document = Document()
    document.add_heading(f"Отчет системы об ошибке", 0)

    p = document.add_paragraph(f'Время инцидента: ')
    p.add_run(f'{date_src}').bold = True

    p = document.add_paragraph(f'Краткое описание ошибки: ')
    p.add_run(f'{error_basic}').bold = True

    document.add_paragraph(f'Полный лог ошибки: \n{error}')

    document.add_paragraph(
        'С уважением, система информирования об ошибках\n'
        'RFC Informer Bot\n'
        'https://t.me/rfc_informer_bot'
    ).italic = True

    document.add_page_break()

    document.add_picture(r'error\tree_error.png', width=Inches(7))
    document.add_picture(r'error\read_error.png', width=Inches(7))

    # document.add_picture(r 'icon\icon.ico', width=Inches(1.25))

    document.save(configs.log_error_filename)


def error_bot(error_basic, us_id_for_error, protect_content_check, client):
    cursor, conn = postgres_init.postgres_init()
    alarm = False
    # current_datetime = datetime.now()
    logs.log_pass(us_id_for_error, 'Ошибка', f'{error_basic}')
    # экранирование спец символов
    e1 = str(error_basic)
    e1 = e1.replace('<', '')
    e1 = e1.replace('>', '')
    try:
        with conn:
            with conn.cursor():
                cursor.execute(f'SELECT mode_test FROM mode_test WHERE user_id = "{us_id_for_error}"')
                mode_test_f_error = cursor.fetchone()[0]
    except (Exception,):
        mode_test_f_error = "no"
        pass

    if mode_test_f_error == "no":
        # получаем ошибку
        error = traceback.format_exc()
        # передаем в функцию записи в файл
        create_file_error(error, error_basic)

        for admin_user_id in firewall_mars.level_list_check('Разработчик'):
            # проверяем, включено ли отображение ошибок
            try:
                cursor.execute(f'SELECT error FROM error_table WHERE user_id = {admin_user_id}')
                error_vision = cursor.fetchone()[0]
            except (Exception,):
                error_vision = "yes"
                pass
            # прерываем цикл если в базе стоит значение нет
            if error_vision == 'yes':
                # ПЫТАЕМСЯ ОТПРАВИТЬ КОД НАТИВНО
                code_error = traceback.format_exc()
                code_error = code_error.replace("<", "")
                code_error = code_error.replace(">", "")
                report = code_error
                if len(report) > 4000:
                    for x in range(0, len(report), 4000):
                        client.send_chat_action(admin_user_id, action="typing")
                        client.send_message(admin_user_id,
                                            f'<pre><code class="language-python">{report[x:x + 4000]}</code></pre>',
                                            parse_mode="html", protect_content=protect_content_check)
                else:
                    client.send_chat_action(admin_user_id, action="typing")
                    client.send_message(admin_user_id,
                                        f'<pre><code class="language-python">{report}</code></pre>',
                                        parse_mode="html", protect_content=protect_content_check)

                # except (Exception,):
                #     pass

                # цикл перебора для уведомления об ошибке
                text1 = "<b>⚠ Внимание! " + \
                        "Произошла ошибка в работе системы!</b>\n" + \
                        " \n" + \
                        "Лог ошибки: <b>" + str(e1) + " </b>\n" + \
                        " \n"
                text2 = "Перезапуск произошел успешно, " + \
                        "система работает в штатном режиме"
                # try:
                text = str(text1) + str(text2)
                client.send_document(admin_user_id,
                                     open(configs.log_error_filename, "rb"),
                                     caption=text,
                                     parse_mode="html")
                alarm = True
                # except (Exception,):
                #     try:
                #         text = str(text1) + \
                #                "Полный лог ошибки, сформировать не удалось\n \n" + \
                #                str(text2)
                #         client.send_message(admin_user_id, text, parse_mode="html",
                #                             protect_content=protect_content_check)
                #         alarm = True
                #     except (Exception,):
                #         pass
                time.sleep(1)
            else:
                pass
    else:
        alarm = False
        pass
        # os.remove("Full error log.doc")
        # если сообщение отправлялось хоть одному админу, указываем это в логах
    if alarm is True:
        info_search = "Администраторы уведомлены об ошибке"
    else:
        info_search = "Администраторы не были уведомлены об ошибке"

    logs.log_pass(us_id_for_error, 'Действие', f'{info_search}')

    # пытаемся уведомить пользователя
    # print(us_id)
    # try:
    #    client.send_message(us_id,
    #                        "⚠ Упс! В моей работе произошла ошибка!\n"+\
    #                        " \n"+\
    #                        "📝 Лог ошибки: <b>"+str(e1)+" </b>\n"+\
    #                        " \n"+\
    #                        "<b>Просьба уведомить администраторов "+\
    #                        "или повторить запрос позже</b>",
    #                        parse_mode="html",
    #                        protect_content=protect_content_check)

    #    full_name = firewall_mars.id_in_name(us_id)
    #    info_search = {"Время":str(current_datetime), \
    #    "Пользователь":"I.R.A.","ID": 78787878,
    #    "Действие":"Пользователь "+full_name+" уведомлен об ошибке"}
    #    print(info_search)
    #    log_pass(info_search, us_id)
    # except:
    #    info_search = {"Время":str(current_datetime), \
    #    "Пользователь":"I.R.A.","ID": 78787878,
    #    "Действие":"Пользователь не уведомлен об ошибке"}
    #    print(info_search)
    #    log_pass(info_search, us_id)
    #    pass
